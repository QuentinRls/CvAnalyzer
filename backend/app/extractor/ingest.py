import pdfplumber
from docx import Document
import chardet
from pathlib import Path
from typing import Union, BinaryIO
from ..utils import logger, CVExtractionError


def read_cv(file: Union[str, Path, BinaryIO]) -> str:
    """
    Extract text from CV file (PDF, DOCX, or TXT)
    
    Args:
        file: File path or file-like object
        
    Returns:
        Extracted text content
        
    Raises:
        CVExtractionError: If file cannot be read or format not supported
    """
    try:
        if isinstance(file, (str, Path)):
            file_path = Path(file)
            if not file_path.exists():
                raise CVExtractionError(f"File not found: {file_path}")
            
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return _read_pdf(file_path)
            elif extension == '.docx':
                return _read_docx(file_path)
            elif extension == '.txt':
                return _read_txt(file_path)
            else:
                raise CVExtractionError(f"Unsupported file format: {extension}")
        else:
            # File-like object - try to detect format from content
            content = file.read()
            file.seek(0)  # Reset for potential re-reading
            
            # Simple format detection based on file signature
            if content.startswith(b'%PDF'):
                return _read_pdf_bytes(content)
            elif content.startswith(b'PK'):  # ZIP signature (DOCX)
                file.seek(0)
                return _read_docx_filelike(file)
            else:
                # Assume text
                return _read_txt_bytes(content)
                
    except Exception as e:
        logger.error(f"Failed to read CV file: {e}")
        raise CVExtractionError(f"Cannot extract text from file: {str(e)}")


def _read_pdf(file_path: Path) -> str:
    """Extract text from PDF file"""
    text_parts = []
    
    with pdfplumber.open(file_path) as pdf:
        if not pdf.pages:
            raise CVExtractionError("PDF contains no pages")
            
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    
    full_text = '\n'.join(text_parts).strip()
    
    if len(full_text) < 50:
        raise CVExtractionError("PDF contains insufficient text (possible scanned document)")
    
    logger.info(f"Extracted {len(full_text)} characters from PDF")
    return full_text


def _read_pdf_bytes(content: bytes) -> str:
    """Extract text from PDF bytes"""
    import io
    text_parts = []
    
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            logger.info(f"PDF opened successfully, number of pages: {len(pdf.pages) if pdf.pages else 0}")
            
            if not pdf.pages:
                raise CVExtractionError("PDF contains no pages")
            
            for i, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    logger.info(f"Page {i+1}: extracted {len(page_text) if page_text else 0} characters")
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i+1}: {e}")
                    continue
        
        full_text = '\n'.join(text_parts).strip()
        
        if len(full_text) < 50:
            raise CVExtractionError("PDF contains insufficient text (possible scanned document)")
        
        return full_text
        
    except Exception as e:
        logger.error(f"Failed to process PDF: {e}")
        # Try alternative PDF reading method
        try:
            import PyPDF2
            logger.info("Trying alternative PDF reading with PyPDF2")
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            if len(pdf_reader.pages) == 0:
                raise CVExtractionError("PDF contains no readable pages")
                
            text_parts = []
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as page_error:
                    logger.warning(f"PyPDF2: Failed to extract text from page {i+1}: {page_error}")
                    continue
            
            full_text = '\n'.join(text_parts).strip()
            if len(full_text) < 50:
                raise CVExtractionError("PDF appears to be scanned or contains no extractable text")
                
            return full_text
            
        except ImportError:
            logger.warning("PyPDF2 not available as fallback")
            raise CVExtractionError(f"Cannot extract text from PDF: {str(e)}")
        except Exception as fallback_error:
            logger.error(f"Both pdfplumber and PyPDF2 failed: {fallback_error}")
            raise CVExtractionError(f"Cannot extract text from PDF: {str(e)}")


def _read_docx(file_path: Path) -> str:
    """Extract text from DOCX file"""
    doc = Document(file_path)
    
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())
    
    full_text = '\n'.join(text_parts).strip()
    
    if len(full_text) < 50:
        raise CVExtractionError("DOCX contains insufficient text")
    
    logger.info(f"Extracted {len(full_text)} characters from DOCX")
    return full_text


def _read_docx_filelike(file: BinaryIO) -> str:
    """Extract text from DOCX file-like object"""
    doc = Document(file)
    
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())
    
    full_text = '\n'.join(text_parts).strip()
    
    if len(full_text) < 50:
        raise CVExtractionError("DOCX contains insufficient text")
    
    return full_text


def _read_txt(file_path: Path) -> str:
    """Extract text from TXT file with encoding detection"""
    # Detect encoding
    with open(file_path, 'rb') as f:
        raw_data = f.read()
    
    encoding_result = chardet.detect(raw_data)
    encoding = encoding_result.get('encoding', 'utf-8')
    
    if encoding is None:
        encoding = 'utf-8'
    
    try:
        text = raw_data.decode(encoding)
    except UnicodeDecodeError:
        # Fallback to utf-8 with error handling
        text = raw_data.decode('utf-8', errors='replace')
    
    text = text.strip()
    
    if len(text) < 50:
        raise CVExtractionError("Text file contains insufficient content")
    
    logger.info(f"Extracted {len(text)} characters from TXT (encoding: {encoding})")
    return text


def _read_txt_bytes(content: bytes) -> str:
    """Extract text from text bytes with encoding detection"""
    encoding_result = chardet.detect(content)
    encoding = encoding_result.get('encoding', 'utf-8')
    
    if encoding is None:
        encoding = 'utf-8'
    
    try:
        text = content.decode(encoding)
    except UnicodeDecodeError:
        text = content.decode('utf-8', errors='replace')
    
    text = text.strip()
    
    if len(text) < 50:
        raise CVExtractionError("Text content insufficient")
    
    return text
